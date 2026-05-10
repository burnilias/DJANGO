import json
from django.http import JsonResponse, HttpResponseRedirect, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Q
from rest_framework.authtoken.models import Token
from django.contrib.auth.hashers import make_password, check_password
from .models import CustomUser, Classe, Groupe, ChatMessage, ChatBan, FriendRequest, Friendship, BlockedUser, PrivateMessage
from .decorators import login_required_custom, role_required, _get_token_user, session_login_required, session_role_required


# ── Helpers ─────────────────────────────────────────────────────────────────

def _user_to_dict(user):
    """Serialise a CustomUser the way the frontend expects."""
    return {
        'id': user.id,
        'email': user.email,
        'first_name': user.first_name,
        'last_name': user.last_name,
        'name': user.name,
        'role': user.role,
        'status': user.status,
        'is_active': user.status == 'active',
        'phone': user.phone,
        'bio': user.bio,
        'created_at': user.created_at.isoformat() if user.created_at else None,
    }


def _parse_json(request):
    try:
        return json.loads(request.body)
    except Exception:
        return {}


# ── Login ──────────────────────────────────────────────────────────────────

@csrf_exempt
@require_http_methods(['POST'])
def login_view(request):
    data = _parse_json(request)
    email = (data.get('email') or '').strip()
    password = data.get('password') or ''

    if not email or not password:
        return JsonResponse(
            {'success': False, 'message': 'Email et mot de passe requis.'},
            status=400,
        )

    try:
        user = CustomUser.objects.get(email=email)
    except CustomUser.DoesNotExist:
        return JsonResponse(
            {'success': False, 'message': 'Email ou mot de passe incorrect.'},
            status=401,
        )

    if not user.check_password(password):
        return JsonResponse(
            {'success': False, 'message': 'Email ou mot de passe incorrect.'},
            status=401,
        )

    if user.status == 'desactive':
        return JsonResponse(
            {'success': False, 'message': 'Ce compte est désactivé.'},
            status=403,
        )

    if user.status == 'veille':
        return JsonResponse(
            {'success': False, 'message': 'Ce compte est en veille.'},
            status=403,
        )

    # Create / retrieve DRF Token
    token, _ = Token.objects.get_or_create(user=user)

    # Also store in Django session (spec requirement)
    request.session['user_id'] = user.id
    request.session['role'] = user.role
    request.session['name'] = user.name
    request.session['first_name'] = user.first_name
    request.session['last_name'] = user.last_name

    return JsonResponse({
        'success': True,
        'token': token.key,
        'user': _user_to_dict(user),
        'role': user.role,
        'message': 'Authentication successful',
    })


# ── Logout ──────────────────────────────────────────────────────────────────

@csrf_exempt
@require_http_methods(['POST'])
def logout_view(request):
    user = _get_token_user(request)
    if user:
        try:
            Token.objects.filter(user=user).delete()
        except Exception:
            pass
    request.session.flush()
    return JsonResponse({'success': True, 'message': 'Logged out successfully.'})


def logout_template_view(request):
    """Template view: logout and redirect to login page"""
    request.session.flush()
    return HttpResponseRedirect('/')


# ── Profile ─────────────────────────────────────────────────────────────────

@login_required_custom
@require_http_methods(['GET'])
def profile_view(request):
    user = request.custom_user
    return JsonResponse({
        'success': True,
        'user': _user_to_dict(user),
        'role': user.role,
    })


# ── Admin: List Users ─────────────────────────────────────────────────────

@login_required_custom
@role_required('admin')
@require_http_methods(['GET'])
def admin_users_list(request):
    users = CustomUser.objects.all().order_by('-created_at')
    return JsonResponse({
        'success': True,
        'users': [_user_to_dict(u) for u in users],
    })


# ── Admin: Create User ─────────────────────────────────────────────────────

@csrf_exempt
@login_required_custom
@role_required('admin')
@require_http_methods(['POST'])
def admin_user_create(request):
    data = _parse_json(request)
    email = (data.get('email') or '').strip()
    password = data.get('password') or ''
    first_name = (data.get('first_name') or '').strip()
    last_name = (data.get('last_name') or '').strip()
    role = data.get('role', 'student')
    phone = data.get('phone', '')
    bio = data.get('bio', '')

    if not email or not password or not first_name or not last_name:
        return JsonResponse(
            {'success': False, 'message': 'Email, password, first name and last name are required.'},
            status=400,
        )

    if role not in ('admin', 'teacher', 'student'):
        return JsonResponse(
            {'success': False, 'message': 'Invalid role.'},
            status=400,
        )

    if CustomUser.objects.filter(email=email).exists():
        return JsonResponse(
            {'success': False, 'message': 'A user with this email already exists.'},
            status=400,
        )

    user = CustomUser(
        email=email,
        first_name=first_name,
        last_name=last_name,
        role=role,
        phone=phone,
        bio=bio,
        status='active',
    )
    user.set_password(password)
    user.save()

    return JsonResponse({
        'success': True,
        'message': 'User created successfully.',
        'user': _user_to_dict(user),
    })


# ── Admin: Update User ─────────────────────────────────────────────────────

@csrf_exempt
@login_required_custom
@role_required('admin')
@require_http_methods(['PUT', 'PATCH'])
def admin_user_update(request, user_id):
    try:
        target = CustomUser.objects.get(pk=user_id)
    except CustomUser.DoesNotExist:
        return JsonResponse(
            {'success': False, 'message': 'User not found.'},
            status=404,
        )

    admin = request.custom_user

    # Safety: admin cannot change their own role or status
    if admin.id == target.id:
        data = _parse_json(request)
        if 'role' in data and data['role'] != admin.role:
            return JsonResponse(
                {'success': False, 'message': 'You cannot change your own role.'},
                status=403,
            )
        # Map is_active back to status
        new_status = _map_is_active_to_status(data)
        if new_status and new_status != admin.status:
            return JsonResponse(
                {'success': False, 'message': 'You cannot change your own status.'},
                status=403,
            )

    data = _parse_json(request)

    if 'email' in data:
        target.email = data['email'].strip()
    if 'first_name' in data:
        target.first_name = data['first_name'].strip()
    if 'last_name' in data:
        target.last_name = data['last_name'].strip()
    if 'role' in data:
        if data['role'] not in ('admin', 'teacher', 'student'):
            return JsonResponse(
                {'success': False, 'message': 'Invalid role.'},
                status=400,
            )
        target.role = data['role']
    if 'phone' in data:
        target.phone = data['phone']
    if 'bio' in data:
        target.bio = data['bio']

    # Map is_active (boolean from frontend) → status field
    new_status = _map_is_active_to_status(data)
    if new_status:
        target.status = new_status

    # Optional password change
    if data.get('password'):
        target.set_password(data['password'])

    target.save()

    return JsonResponse({
        'success': True,
        'message': 'User updated successfully.',
        'user': _user_to_dict(target),
    })


def _map_is_active_to_status(data):
    """Convert frontend's is_active boolean to our status string."""
    if 'is_active' in data:
        return 'active' if data['is_active'] else 'desactive'
    if 'status' in data:
        return data['status']
    return None


# ── Admin: Delete User ─────────────────────────────────────────────────────

@csrf_exempt
@login_required_custom
@role_required('admin')
@require_http_methods(['DELETE'])
def admin_user_delete(request, user_id):
    admin = request.custom_user

    if admin.id == int(user_id):
        return JsonResponse(
            {'success': False, 'message': 'You cannot delete your own account.'},
            status=403,
        )

    try:
        target = CustomUser.objects.get(pk=user_id)
    except CustomUser.DoesNotExist:
        return JsonResponse(
            {'success': False, 'message': 'User not found.'},
            status=404,
        )

    target.delete()
    return JsonResponse({
        'success': True,
        'message': 'User deleted successfully.',
    })


# ══════════════════════════════════════════════════════════════════════════
# ADMIN DASHBOARD — Django Template Views (session-based auth)
# ══════════════════════════════════════════════════════════════════════════

@session_login_required
@session_role_required('admin')
def admin_dashboard(request):
    user = request.session_user
    context = {
        'total_users': CustomUser.objects.count(),
        'total_admins': CustomUser.objects.filter(role='admin').count(),
        'total_teachers': CustomUser.objects.filter(role='teacher').count(),
        'total_students': CustomUser.objects.filter(role='student').count(),
        'recent_users': CustomUser.objects.order_by('-created_at')[:5],
        'active_count': CustomUser.objects.filter(status='active').count(),
        'desactive_count': CustomUser.objects.filter(status='desactive').count(),
        'veille_count': CustomUser.objects.filter(status='veille').count(),
        'session_user': user,
        'page_title': 'Dashboard',
    }
    return render(request, 'admin/dashboard.html', context)


@session_login_required
@session_role_required('admin')
def admin_user_list(request):
    user = request.session_user
    search = request.GET.get('search', '')
    role_filter = request.GET.get('role', '')
    status_filter = request.GET.get('status', '')

    qs = CustomUser.objects.all()
    if search:
        qs = qs.filter(
            Q(first_name__icontains=search) |
            Q(last_name__icontains=search) |
            Q(email__icontains=search)
        )
    if role_filter:
        qs = qs.filter(role=role_filter)
    if status_filter:
        qs = qs.filter(status=status_filter)
    qs = qs.order_by('-created_at')

    paginator = Paginator(qs, 10)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    context = {
        'page_obj': page_obj,
        'search': search,
        'role_filter': role_filter,
        'status_filter': status_filter,
        'session_user': user,
        'page_title': 'Utilisateurs',
        'groupes': Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom'),
    }
    return render(request, 'admin/users/list.html', context)


@session_login_required
@session_role_required('admin')
def admin_add_user(request):
    if request.method == 'POST':
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        email = request.POST.get('email', '').strip()
        phone = request.POST.get('phone', '').strip()
        bio = request.POST.get('bio', '').strip()
        role = request.POST.get('role', 'student')
        status = request.POST.get('status', 'active')
        password = request.POST.get('password', '')
        password_confirm = request.POST.get('password_confirm', '')

        errors = {}
        if not first_name:
            errors['first_name'] = 'Le prénom est requis.'
        if not last_name:
            errors['last_name'] = 'Le nom est requis.'
        if not email:
            errors['email'] = "L'email est requis."
        elif CustomUser.objects.filter(email=email).exists():
            errors['email'] = 'Un utilisateur avec cet email existe déjà.'
        if not password:
            errors['password'] = 'Le mot de passe est requis.'
        elif len(password) < 8:
            errors['password'] = 'Le mot de passe doit contenir au moins 8 caractères.'
        if password != password_confirm:
            errors['password_confirm'] = 'Les mots de passe ne correspondent pas.'
        if role not in ('admin', 'teacher', 'student'):
            errors['role'] = 'Rôle invalide.'
        if status not in ('active', 'desactive', 'veille'):
            errors['status'] = 'Statut invalide.'

        if errors:
            # Return JSON for AJAX requests
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'errors': errors}, status=400)
            return render(request, 'admin/users/add.html', {
                'errors': errors,
                'form_data': request.POST,
                'session_user': request.session_user,
                'page_title': 'Ajouter un utilisateur',
            })

        groupe_id = request.POST.get('groupe_id', '') or None
        new_user = CustomUser(
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone,
            bio=bio,
            role=role,
            status=status,
            groupe_id=groupe_id,
        )
        new_user.set_password(password)
        new_user.save()
        
        # Return JSON for AJAX requests
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'message': 'Utilisateur créé avec succès.'})
        
        messages.success(request, 'Utilisateur créé avec succès.')
        return redirect('admin_user_list')

    # GET request - redirect to user list (modal handles the form now)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'success': False, 'message': 'Method not allowed.'}, status=405)
    
    return render(request, 'admin/users/add.html', {
        'session_user': request.session_user,
        'page_title': 'Ajouter un utilisateur',
    })


@session_login_required
@session_role_required('admin')
def admin_edit_user(request, id):
    target = get_object_or_404(CustomUser, id=id)
    admin_user = request.session_user

    if request.method == 'GET':
        return JsonResponse({
            'id': target.id,
            'first_name': target.first_name,
            'last_name': target.last_name,
            'email': target.email,
            'phone': target.phone,
            'bio': target.bio,
            'role': target.role,
            'status': target.status,
            'groupe_id': target.groupe_id,
            'is_self': admin_user.id == target.id,
        })

    # POST
    first_name = request.POST.get('first_name', '').strip()
    last_name = request.POST.get('last_name', '').strip()
    email = request.POST.get('email', '').strip()
    phone = request.POST.get('phone', '').strip()
    bio = request.POST.get('bio', '').strip()
    role = request.POST.get('role', target.role)
    status = request.POST.get('status', target.status)
    groupe_id = request.POST.get('groupe_id', '') or None
    new_password = request.POST.get('new_password', '')
    new_password_confirm = request.POST.get('new_password_confirm', '')

    errors = {}
    if not first_name:
        errors['first_name'] = 'Le prénom est requis.'
    if not last_name:
        errors['last_name'] = 'Le nom est requis.'
    if not email:
        errors['email'] = "L'email est requis."
    elif CustomUser.objects.filter(email=email).exclude(id=target.id).exists():
        errors['email'] = 'Un utilisateur avec cet email existe déjà.'

    # Admin cannot change own role/status
    is_self = admin_user.id == target.id
    if is_self:
        role = target.role
        status = target.status
    else:
        if role not in ('admin', 'teacher', 'student'):
            errors['role'] = 'Rôle invalide.'
        if status not in ('active', 'desactive', 'veille'):
            errors['status'] = 'Statut invalide.'

    if new_password:
        if len(new_password) < 8:
            errors['new_password'] = 'Le mot de passe doit contenir au moins 8 caractères.'
        if new_password != new_password_confirm:
            errors['new_password_confirm'] = 'Les mots de passe ne correspondent pas.'

    if errors:
        return JsonResponse({'success': False, 'errors': errors}, status=400)

    target.first_name = first_name
    target.last_name = last_name
    target.email = email
    target.phone = phone
    target.bio = bio
    target.role = role
    target.status = status
    target.groupe_id = groupe_id
    if new_password:
        target.set_password(new_password)
    target.save()

    return JsonResponse({'success': True, 'message': 'Utilisateur mis à jour.'})


@session_login_required
@session_role_required('admin')
def admin_delete_user(request, id):
    admin_user = request.session_user
    if admin_user.id == id:
        messages.error(request, 'Vous ne pouvez pas supprimer votre propre compte.')
        return redirect('admin_user_list')

    target = get_object_or_404(CustomUser, id=id)
    # Delete token first to avoid FK constraint error
    Token.objects.filter(user=target).delete()
    target.delete()
    messages.success(request, 'Utilisateur supprimé.')
    return redirect('admin_user_list')


@session_login_required
@session_role_required('admin')
def admin_change_status(request, id):
    admin_user = request.session_user
    if admin_user.id == id:
        return JsonResponse({'success': False, 'error': 'Vous ne pouvez pas modifier votre propre statut.'}, status=403)

    target = get_object_or_404(CustomUser, id=id)
    if request.method == 'POST':
        data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
        new_status = data.get('new_status', '')
        if new_status not in ('active', 'desactive', 'veille'):
            return JsonResponse({'success': False, 'error': 'Statut invalide.'}, status=400)
        target.status = new_status
        target.save()
        return JsonResponse({'success': True, 'new_status': new_status})
    return JsonResponse({'success': False, 'error': 'Méthode non autorisée.'}, status=405)


@session_login_required
@session_role_required('admin')
def admin_profile(request):
    user = request.session_user
    # Refresh from DB
    user = CustomUser.objects.get(pk=user.id)

    if request.method == 'POST':
        form_type = request.POST.get('form_type', '')

        if form_type == 'update_info':
            first_name = request.POST.get('first_name', '').strip()
            last_name = request.POST.get('last_name', '').strip()
            phone = request.POST.get('phone', '').strip()
            bio = request.POST.get('bio', '').strip()
            if first_name:
                user.first_name = first_name
            if last_name:
                user.last_name = last_name
            user.phone = phone
            user.bio = bio
            user.save()
            # Update session
            request.session['first_name'] = user.first_name
            request.session['last_name'] = user.last_name
            messages.success(request, 'Profil mis à jour.')
            return redirect('admin_profile')

        elif form_type == 'change_password':
            current_password = request.POST.get('current_password', '')
            new_password = request.POST.get('new_password', '')
            new_password_confirm = request.POST.get('new_password_confirm', '')

            if not user.check_password(current_password):
                messages.error(request, 'Mot de passe actuel incorrect.')
            elif len(new_password) < 8:
                messages.error(request, 'Le nouveau mot de passe doit contenir au moins 8 caractères.')
            elif new_password != new_password_confirm:
                messages.error(request, 'Les mots de passe ne correspondent pas.')
            else:
                user.set_password(new_password)
                user.save()
                # Re-create session to avoid session invalidation
                request.session['user_id'] = user.id
                request.session['role'] = user.role
                request.session['first_name'] = user.first_name
                request.session['last_name'] = user.last_name
                messages.success(request, 'Mot de passe mis à jour.')
            return redirect('admin_profile')

    context = {
        'profile_user': user,
        'session_user': user,
        'page_title': 'Mon Profil',
    }
    return render(request, 'admin/profile.html', context)


@session_login_required
@session_role_required('admin')
def admin_settings(request):
    context = {
        'session_user': request.session_user,
        'page_title': 'Paramètres',
    }
    return render(request, 'admin/settings.html', context)


# ══════════════════════════════════════════════════════════════════════════
# EMPLOI DU TEMPS (Schedule)
# ══════════════════════════════════════════════════════════════════════════

@session_login_required
@session_role_required('admin')
def admin_schedule_list(request):
    from .models import Schedule
    from datetime import datetime, timedelta
    user = request.session_user
    
    # AJAX request for group schedule
    groupe_id = request.GET.get('groupe_id')
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    if groupe_id and is_ajax:
        schedules = Schedule.objects.select_related('teacher', 'groupe').filter(groupe_id=groupe_id)
        data = []
        for s in schedules:
            data.append({
                'id': s.id,
                'name': s.name,
                'room': s.room,
                'day_of_week': s.day_of_week,
                'start_time': s.start_time.strftime('%H:%M'),
                'end_time': s.end_time.strftime('%H:%M'),
                'schedule_type': s.schedule_type,
            })
        return JsonResponse({'success': True, 'schedules': data})

    # Full page render — only show schedules when a group is selected
    selected_groupe_id = groupe_id
    if groupe_id:
        all_schedules = Schedule.objects.select_related('teacher').filter(groupe_id=groupe_id)
    else:
        all_schedules = Schedule.objects.none()
    teachers = CustomUser.objects.filter(role='teacher', status='active')
    groupes = Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom')
    
    # Get current week dates
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    week_start = monday.strftime('%d %B')
    week_end = (monday + timedelta(days=6)).strftime('%d %B %Y')
    
    # Day mapping
    day_names_fr = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']
    day_short_fr = ['Lun', 'Mar', 'Mer', 'Jeu', 'Ven', 'Sam', 'Dim']
    today_weekday = today.weekday()  # 0=Monday
    
    # Build day headers
    day_names = []
    for i in range(7):
        date = monday + timedelta(days=i)
        day_names.append({
            'name': day_names_fr[i],
            'short': day_short_fr[i],
            'date': date.strftime('%d'),
            'num': str(i + 1),
            'today': i == today_weekday
        })
    
    # Time slots (8:00 to 20:00, every hour) - 60px per hour = 720px total
    time_slots = []
    for i in range(13):  # 8:00 to 20:00
        hour = 8 + i
        time_slots.append({
            'label': f'{hour:02d}:00',
            'position': i * 60  # 60px per hour
        })
    
    # Build days with their schedules positioned correctly
    days = []
    for day_num in range(1, 8):
        day_schedules = []
        day_is_today = (day_num - 1) == today_weekday
        
        for schedule in all_schedules:
            if int(schedule.day_of_week) == day_num:
                # Calculate position based on time (60px per hour, starting at 8:00)
                start_h = schedule.start_time.hour
                start_m = schedule.start_time.minute
                end_h = schedule.end_time.hour
                end_m = schedule.end_time.minute
                
                # Position from 8:00 (minute 0)
                start_minutes = (start_h - 8) * 60 + start_m
                end_minutes = (end_h - 8) * 60 + end_m
                duration = end_minutes - start_minutes
                
                # Map schedule types to user's CSS classes
                type_class_map = {
                    'cours': 'type-lab',
                    'controle': 'type-lang',
                    'examen': 'type-research',
                    'rattrapage': 'type-research',  # fallback
                }
                type_class = type_class_map.get(schedule.schedule_type, '')
                
                day_schedules.append({
                    'id': schedule.id,
                    'name': schedule.name,
                    'start_time': schedule.start_time,
                    'end_time': schedule.end_time,
                    'room': schedule.room,
                    'teacher': schedule.teacher,
                    'type_label': schedule.get_schedule_type_display() if hasattr(schedule, 'get_schedule_type_display') else schedule.schedule_type,
                    'top': start_minutes,
                    'height': duration,
                    'type_class': type_class,
                })
        
        days.append({
            'num': str(day_num),
            'today': day_is_today,
            'schedules': day_schedules
        })
    
    context = {
        'day_names': day_names,
        'days': days,
        'time_slots': time_slots,
        'teachers': teachers,
        'groupes': groupes,
        'selected_groupe_id': selected_groupe_id,
        'week_start': week_start,
        'week_end': week_end,
        'session_user': user,
        'page_title': 'Emploi du temps',
    }
    return render(request, 'admin/emploi/list.html', context)


@session_login_required
@session_role_required('admin')
def admin_schedule_add(request):
    from .models import Schedule
    user = request.session_user
    
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        room = request.POST.get('room', '').strip()
        day_of_week = request.POST.get('day_of_week', '1')
        start_time = request.POST.get('start_time', '')
        end_time = request.POST.get('end_time', '')
        schedule_type = request.POST.get('schedule_type', 'cours')
        teacher_id = request.POST.get('teacher', '')
        groupe_id = request.POST.get('groupe', '')
        
        errors = {}
        if not name:
            errors['name'] = 'Le nom est requis.'
        if not start_time:
            errors['start_time'] = "L'heure de début est requise."
        if not end_time:
            errors['end_time'] = "L'heure de fin est requise."
        if not groupe_id:
            errors['groupe'] = 'Le groupe est requis.'
        
        if errors:
            teachers = CustomUser.objects.filter(role='teacher', status='active')
            groupes = Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom')
            return render(request, 'admin/emploi/add.html', {
                'errors': errors,
                'form_data': request.POST,
                'teachers': teachers,
                'groupes': groupes,
                'session_user': user,
                'page_title': 'Ajouter un cours',
            })
        
        schedule = Schedule(
            name=name,
            description=description,
            room=room,
            day_of_week=day_of_week,
            start_time=start_time,
            end_time=end_time,
            schedule_type=schedule_type,
            teacher_id=teacher_id if teacher_id else None,
            groupe_id=groupe_id if groupe_id else None,
        )
        schedule.save()
        messages.success(request, 'Cours ajouté avec succès.')
        return redirect('admin_schedule_list')
    
    teachers = CustomUser.objects.filter(role='teacher', status='active')
    groupes = Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom')
    return render(request, 'admin/emploi/add.html', {
        'teachers': teachers,
        'groupes': groupes,
        'session_user': user,
        'page_title': 'Ajouter un cours',
    })


@session_login_required
@session_role_required('admin')
def admin_schedule_edit(request, id):
    from .models import Schedule
    schedule = get_object_or_404(Schedule, id=id)
    user = request.session_user
    
    if request.method == 'POST':
        schedule.name = request.POST.get('name', '').strip()
        schedule.description = request.POST.get('description', '').strip()
        schedule.room = request.POST.get('room', '').strip()
        schedule.day_of_week = request.POST.get('day_of_week', '1')
        schedule.start_time = request.POST.get('start_time', '')
        schedule.end_time = request.POST.get('end_time', '')
        schedule.schedule_type = request.POST.get('schedule_type', 'cours')
        teacher_id = request.POST.get('teacher', '')
        schedule.teacher_id = teacher_id if teacher_id else None
        groupe_id = request.POST.get('groupe', '')
        schedule.groupe_id = groupe_id if groupe_id else None
        schedule.save()
        messages.success(request, 'Cours mis à jour.')
        return redirect('admin_schedule_list')
    
    teachers = CustomUser.objects.filter(role='teacher', status='active')
    groupes = Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom')
    return render(request, 'admin/emploi/edit.html', {
        'schedule': schedule,
        'teachers': teachers,
        'groupes': groupes,
        'session_user': user,
        'page_title': 'Modifier le cours',
    })


@session_login_required
@session_role_required('admin')
def admin_schedule_delete(request, id):
    from .models import Schedule
    schedule = get_object_or_404(Schedule, id=id)
    schedule.delete()
    messages.success(request, 'Cours supprimé.')
    return redirect('admin_schedule_list')


# ══════════════════════════════════════════════════════════════════════════
# STUDENT VIEWS (Readonly)
# ══════════════════════════════════════════════════════════════════════════

@session_login_required
@session_role_required('student')
def student_dashboard(request):
    user = request.session_user
    return render(request, 'student/dashboard.html', {
        'session_user': user,
        'page_title': 'Dashboard Étudiant',
    })


@session_login_required
@session_role_required('student')
def student_schedule(request):
    from .models import Schedule
    from datetime import datetime, timedelta
    user = request.session_user
    
    # Filter schedules by student's group — null-group entries never shown to students
    user_group = user.groupe
    if user_group:
        all_schedules = Schedule.objects.select_related('teacher').filter(groupe=user_group)
    else:
        all_schedules = Schedule.objects.none()
    
    # Get current week dates
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    week_start = monday.strftime('%d %B')
    week_end = (monday + timedelta(days=6)).strftime('%d %B %Y')
    
    # Day mapping
    day_names_fr = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']
    day_short_fr = ['Lun', 'Mar', 'Mer', 'Jeu', 'Ven', 'Sam', 'Dim']
    today_weekday = today.weekday()  # 0=Monday
    
    # Build day headers
    day_names = []
    for i in range(7):
        date = monday + timedelta(days=i)
        day_names.append({
            'name': day_names_fr[i],
            'short': day_short_fr[i],
            'date': date.strftime('%d'),
            'num': str(i + 1),
            'today': i == today_weekday
        })
    
    # Time slots
    time_slots = []
    for i in range(13):
        hour = 8 + i
        time_slots.append({
            'label': f'{hour:02d}:00',
            'position': i * 60
        })
    
    # Build days with schedules
    days = []
    for day_num in range(1, 8):
        day_schedules = []
        day_is_today = (day_num - 1) == today_weekday
        
        for schedule in all_schedules:
            if int(schedule.day_of_week) == day_num:
                start_h = schedule.start_time.hour
                start_m = schedule.start_time.minute
                end_h = schedule.end_time.hour
                end_m = schedule.end_time.minute
                
                start_minutes = (start_h - 8) * 60 + start_m
                end_minutes = (end_h - 8) * 60 + end_m
                duration = end_minutes - start_minutes
                
                # Map schedule types to user's CSS classes
                type_class_map = {
                    'cours': 'type-lab',
                    'controle': 'type-lang',
                    'examen': 'type-research',
                    'rattrapage': 'type-rattrapage',
                }
                type_class = type_class_map.get(schedule.schedule_type, '')
                
                day_schedules.append({
                    'id': schedule.id,
                    'name': schedule.name,
                    'start_time': schedule.start_time,
                    'end_time': schedule.end_time,
                    'room': schedule.room,
                    'teacher': schedule.teacher,
                    'type_label': schedule.get_schedule_type_display() if hasattr(schedule, 'get_schedule_type_display') else schedule.schedule_type,
                    'day_name': day_names_fr[day_num - 1],
                    'top': start_minutes,
                    'height': duration,
                    'type_class': type_class,
                })
        
        days.append({
            'num': str(day_num),
            'today': day_is_today,
            'day_name': day_names_fr[day_num - 1],
            'schedules': day_schedules
        })
    
    context = {
        'day_names': day_names,
        'days': days,
        'time_slots': time_slots,
        'week_start': week_start,
        'week_end': week_end,
        'session_user': user,
        'page_title': 'Mon Emploi du temps',
    }
    return render(request, 'student/emploi.html', context)


@session_login_required
@session_role_required('student')
def student_profile(request):
    user = request.session_user
    return render(request, 'student/profile.html', {
        'session_user': user,
        'page_title': 'Mon Profil',
    })


# ── Classes & Groupes API ──
@session_login_required
@session_role_required('admin')
def api_classes_list(request):
    if request.method == 'POST':
        nom = json.loads(request.body).get('nom') if request.content_type == 'application/json' else request.POST.get('nom')
        if nom not in dict(Classe.NOM_CHOICES):
            return JsonResponse({'success': False, 'message': 'Nom invalide.'}, status=400)
        classe = Classe.objects.create(nom=nom)
        return JsonResponse({'success': True, 'message': 'Classe créée.', 'id': classe.id})
    classes = Classe.objects.all().order_by('nom')
    data = [{'id': c.id, 'nom': c.nom, 'groupes_count': c.groupes.count()} for c in classes]
    return JsonResponse({'success': True, 'classes': data})


@session_login_required
@session_role_required('admin')
def api_groupes_list(request):
    if request.method == 'POST':
        data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
        nom = data.get('nom')
        classe_id = data.get('classe_id')
        if not nom or not classe_id:
            return JsonResponse({'success': False, 'message': 'Nom et classe requis.'}, status=400)
        groupe = Groupe.objects.create(nom=nom, classe_id=classe_id)
        return JsonResponse({'success': True, 'message': 'Groupe créé.', 'id': groupe.id})
    classe_id = request.GET.get('classe_id')
    groupes = Groupe.objects.all()
    if classe_id:
        groupes = groupes.filter(classe_id=classe_id)
    groupes = groupes.order_by('classe__nom', 'nom')
    data = [{'id': g.id, 'nom': g.nom, 'classe': g.classe.nom, 'members_count': g.members.count()} for g in groupes]
    return JsonResponse({'success': True, 'groupes': data})


@session_login_required
@session_role_required('admin')
def api_affecter_groupe(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'POST requis.'}, status=405)
    data = json.loads(request.body) if request.content_type == 'application/json' else request.POST
    user_id = data.get('user_id')
    groupe_id = data.get('groupe_id')
    try:
        user = CustomUser.objects.get(id=user_id)
        if groupe_id:
            groupe = Groupe.objects.get(id=groupe_id)
            user.groupe = groupe
        else:
            user.groupe = None
        user.save()
        return JsonResponse({'success': True, 'message': 'Groupe affecté.'})
    except (CustomUser.DoesNotExist, Groupe.DoesNotExist):
        return JsonResponse({'success': False, 'message': 'Utilisateur ou groupe introuvable.'}, status=404)


@session_login_required
@session_role_required('admin')
def admin_gestion_classes(request):
    classes = Classe.objects.all().order_by('nom')
    groupes = Groupe.objects.select_related('classe').all().order_by('classe__nom', 'nom')
    users = CustomUser.objects.select_related('groupe').filter(role='student').order_by('last_name')
    return render(request, 'admin/classes/gestion_classes.html', {
        'session_user': request.session_user,
        'page_title': 'Classes & Groupes',
        'classes': classes,
        'groupes': groupes,
        'users': users,
    })


# ══════════════════════════════════════════════════════════════════════════
# LMS VIEWS - Admin, Teacher, Student
# ══════════════════════════════════════════════════════════════════════════

from .models import Course, Lesson, Enrollment, Progress, Certificate, StudentOfMonth, LessonVideo, VideoCompletion
from django.utils import timezone
from datetime import datetime, timedelta
import uuid
import os
from django.conf import settings

# ── Admin LMS Management ──────────────────────────────────────────────────

@session_login_required
@session_role_required('admin')
def admin_courses_list(request):
    """Admin can view and manage all courses"""
    search = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    teacher_filter = request.GET.get('teacher', '')
    
    courses = Course.objects.select_related('teacher').all()
    
    if search:
        courses = courses.filter(Q(title__icontains=search) | Q(description__icontains=search))
    if status_filter:
        courses = courses.filter(status=status_filter)
    if teacher_filter:
        courses = courses.filter(teacher_id=teacher_filter)
    
    courses = courses.order_by('-created_at')
    teachers = CustomUser.objects.filter(role='teacher', status='active')
    
    context = {
        'courses': courses,
        'teachers': teachers,
        'search': search,
        'status_filter': status_filter,
        'teacher_filter': teacher_filter,
        'session_user': request.session_user,
        'page_title': 'Gestion des Cours',
    }
    return render(request, 'admin/courses/list.html', context)


@session_login_required
@session_role_required('admin')
def admin_course_delete(request, course_id):
    """Admin can delete any course"""
    course = get_object_or_404(Course, id=course_id)
    course.delete()
    messages.success(request, 'Cours supprimé avec succès.')
    return redirect('admin_courses_list')


@session_login_required
@session_role_required('admin')
def admin_course_edit(request, course_id):
    """Admin can edit any course"""
    course = get_object_or_404(Course, id=course_id)
    teachers = CustomUser.objects.filter(role='teacher', status='active')
    
    if request.method == 'POST':
        course.title = request.POST.get('title', '').strip()
        course.description = request.POST.get('description', '').strip()
        course.status = request.POST.get('status', 'draft')
        teacher_id = request.POST.get('teacher', '')
        if teacher_id:
            course.teacher_id = teacher_id
        if request.FILES.get('image'):
            course.image = request.FILES['image']
        course.save()
        messages.success(request, 'Cours mis à jour.')
        return redirect('admin_courses_list')
    
    context = {
        'course': course,
        'teachers': teachers,
        'session_user': request.session_user,
        'page_title': 'Modifier le Cours',
    }
    return render(request, 'admin/courses/edit.html', context)


@session_login_required
@session_role_required('admin')
def admin_course_lessons(request, course_id):
    """Admin can view and manage lessons for any course"""
    course = get_object_or_404(Course, id=course_id)
    lessons = course.lessons.all()
    
    context = {
        'course': course,
        'lessons': lessons,
        'session_user': request.session_user,
        'page_title': f'Leçons - {course.title}',
    }
    return render(request, 'admin/courses/lessons.html', context)


@session_login_required
@session_role_required('admin')
def admin_course_enrollments(request, course_id):
    """Admin can view enrollments for any course"""
    course = get_object_or_404(Course, id=course_id)
    enrollments = course.enrollments.select_related('student').all()
    
    context = {
        'course': course,
        'enrollments': enrollments,
        'session_user': request.session_user,
        'page_title': f'Inscriptions - {course.title}',
    }
    return render(request, 'admin/courses/enrollments.html', context)


@session_login_required
@session_role_required('admin')
def admin_all_enrollments(request):
    """Admin can view all enrollments"""
    enrollments = Enrollment.objects.select_related('student', 'course').all().order_by('-enrolled_at')
    
    context = {
        'enrollments': enrollments,
        'session_user': request.session_user,
        'page_title': 'Toutes les Inscriptions',
    }
    return render(request, 'admin/courses/all_enrollments.html', context)


@session_login_required
@session_role_required('admin')
def admin_all_certificates(request):
    """Admin can view all certificates issued"""
    certificates = Certificate.objects.select_related('enrollment__student', 'enrollment__course').all().order_by('-issued_at')
    
    context = {
        'certificates': certificates,
        'session_user': request.session_user,
        'page_title': 'Tous les Certificats',
    }
    return render(request, 'admin/courses/all_certificates.html', context)


@session_login_required
def admin_student_of_month(request):
    """Admin can view and manually trigger Student of the Month calculation"""
    current_month = timezone.now().replace(day=1)
    
    # Get current student of the month
    current_som = StudentOfMonth.objects.filter(month=current_month).first()
    
    # Get previous winners
    previous_winners = StudentOfMonth.objects.select_related('student').exclude(month=current_month)[:12]
    
    context = {
        'current_som': current_som,
        'previous_winners': previous_winners,
        'session_user': request.session_user,
        'page_title': 'Étudiant du Mois',
    }
    return render(request, 'admin/courses/student_of_month.html', context)


# ── Teacher LMS Views ─────────────────────────────────────────────────────

@session_login_required
@session_role_required('teacher')
def teacher_dashboard(request):
    """Teacher dashboard with their courses"""
    user = request.session_user
    courses = Course.objects.filter(teacher=user).order_by('-created_at')
    total_students = Enrollment.objects.filter(course__teacher=user).count()
    published_courses = courses.filter(status='published').count()
    
    context = {
        'courses': courses,
        'total_students': total_students,
        'published_courses': published_courses,
        'total_courses': courses.count(),
        'session_user': user,
        'page_title': 'Dashboard Enseignant',
    }
    return render(request, 'teacher/dashboard.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_courses(request):
    """Teacher can view and manage their own courses"""
    user = request.session_user
    courses = Course.objects.filter(teacher=user).order_by('-created_at')
    
    context = {
        'courses': courses,
        'session_user': user,
        'page_title': 'Mes Cours',
    }
    return render(request, 'teacher/courses/list.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_course_create(request):
    """Teacher can create a new course"""
    user = request.session_user
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        status = request.POST.get('status', 'draft')
        
        if not title:
            messages.error(request, 'Le titre est requis.')
            return render(request, 'teacher/courses/create.html', {
                'session_user': user,
                'page_title': 'Créer un Cours',
                'form_data': request.POST,
            })
        
        course = Course(
            title=title,
            description=description,
            teacher=user,
            status=status,
        )
        if request.FILES.get('image'):
            course.image = request.FILES['image']
        course.save()
        messages.success(request, 'Cours créé avec succès.')
        return redirect('teacher_courses')
    
    context = {
        'session_user': user,
        'page_title': 'Créer un Cours',
    }
    return render(request, 'teacher/courses/create.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_course_edit(request, course_id):
    """Teacher can edit their own course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    
    if request.method == 'POST':
        course.title = request.POST.get('title', '').strip()
        course.description = request.POST.get('description', '').strip()
        course.status = request.POST.get('status', 'draft')
        if request.FILES.get('image'):
            course.image = request.FILES['image']
        course.save()
        messages.success(request, 'Cours mis à jour.')
        return redirect('teacher_courses')
    
    context = {
        'course': course,
        'session_user': user,
        'page_title': 'Modifier le Cours',
    }
    return render(request, 'teacher/courses/edit.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_course_delete(request, course_id):
    """Teacher can delete their own course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    course.delete()
    messages.success(request, 'Cours supprimé.')
    return redirect('teacher_courses')


@session_login_required
@session_role_required('teacher')
def teacher_course_lessons(request, course_id):
    """Teacher can manage lessons for their course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    lessons = course.lessons.all()
    
    context = {
        'course': course,
        'lessons': lessons,
        'session_user': user,
        'page_title': f'Leçons - {course.title}',
    }
    return render(request, 'teacher/courses/lessons.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_lesson_create(request, course_id):
    """Teacher can create a lesson for their course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        content = request.POST.get('content', '').strip()
        video_url = request.POST.get('video_url', '').strip()
        order = request.POST.get('order', 0)
        duration_minutes = request.POST.get('duration_minutes', 0)
        
        if not title:
            messages.error(request, 'Le titre est requis.')
            return render(request, 'teacher/lessons/create.html', {
                'course': course,
                'session_user': user,
                'page_title': 'Ajouter une Leçon',
                'form_data': request.POST,
                'groupes': groupes,
            })
        
        lesson = Lesson(
            course=course,
            title=title,
            content=content,
            video_url=video_url,
            order=int(order) if order else 0,
            duration_minutes=int(duration_minutes) if duration_minutes else 0,
        )
        if groupe_id:
            lesson.groupe_id = groupe_id
        lesson.save()

        # Save multiple video links
        video_urls = request.POST.getlist('video_urls[]')
        video_titles = request.POST.getlist('video_titles[]')
        for i, vurl in enumerate(video_urls):
            vurl = vurl.strip()
            if vurl:
                vtitle = video_titles[i].strip() if i < len(video_titles) else ''
                LessonVideo.objects.create(lesson=lesson, title=vtitle, url=vurl, order=i)

        messages.success(request, 'Leçon ajoutée.')
        return redirect('teacher_course_lessons', course_id=course.id)
    
    context = {
        'course': course,
        'session_user': user,
        'page_title': 'Ajouter une Leçon',
    }
    return render(request, 'teacher/lessons/create.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_lesson_edit(request, course_id, lesson_id):
    """Teacher can edit a lesson in their course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    lesson = get_object_or_404(Lesson, id=lesson_id, course=course)

    if request.method == 'POST':
        lesson.title = request.POST.get('title', '').strip()
        lesson.content = request.POST.get('content', '').strip()
        lesson.video_url = request.POST.get('video_url', '').strip()
        lesson.order = int(request.POST.get('order', 0))
        lesson.duration_minutes = int(request.POST.get('duration_minutes', 0))
        lesson.save()

        # Replace all video links
        lesson.videos.all().delete()
        video_urls = request.POST.getlist('video_urls[]')
        video_titles = request.POST.getlist('video_titles[]')
        for i, vurl in enumerate(video_urls):
            vurl = vurl.strip()
            if vurl:
                vtitle = video_titles[i].strip() if i < len(video_titles) else ''
                LessonVideo.objects.create(lesson=lesson, title=vtitle, url=vurl, order=i)

        messages.success(request, 'Leçon mise à jour.')
        return redirect('teacher_course_lessons', course_id=course.id)

    context = {
        'course': course,
        'lesson': lesson,
        'videos': lesson.videos.all(),
        'session_user': user,
        'page_title': 'Modifier la Leçon',
    }
    return render(request, 'teacher/lessons/edit.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_lesson_delete(request, course_id, lesson_id):
    """Teacher can delete a lesson from their course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    lesson = get_object_or_404(Lesson, id=lesson_id, course=course)
    lesson.delete()
    messages.success(request, 'Leçon supprimée.')
    return redirect('teacher_course_lessons', course_id=course.id)


@session_login_required
@session_role_required('teacher')
def teacher_course_enrollments(request, course_id):
    """Teacher can view enrollments for their course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, teacher=user)
    enrollments = course.enrollments.select_related('student').all()
    
    context = {
        'course': course,
        'enrollments': enrollments,
        'session_user': user,
        'page_title': f'Inscriptions - {course.title}',
    }
    return render(request, 'teacher/courses/enrollments.html', context)


@session_login_required
@session_role_required('teacher')
def teacher_profile(request):
    """Teacher profile page"""
    user = request.session_user
    return render(request, 'teacher/profile.html', {
        'session_user': user,
        'page_title': 'Mon Profil',
    })


# ── Student LMS Views ─────────────────────────────────────────────────────

@session_login_required
@session_role_required('student')
def student_courses_browse(request):
    """Student can browse all published courses"""
    user = request.session_user
    search = request.GET.get('search', '')
    
    courses = Course.objects.filter(status='published').select_related('teacher')
    if search:
        courses = courses.filter(Q(title__icontains=search) | Q(description__icontains=search))
    courses = courses.order_by('-created_at')
    
    # Get student's enrolled course IDs
    enrolled_course_ids = Enrollment.objects.filter(student=user).values_list('course_id', flat=True)
    
    context = {
        'courses': courses,
        'enrolled_course_ids': list(enrolled_course_ids),
        'search': search,
        'session_user': user,
        'page_title': 'Parcourir les Cours',
    }
    return render(request, 'student/courses/browse.html', context)


@session_login_required
@session_role_required('student')
def student_course_detail(request, course_id):
    """Student can view course details and enroll"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, status='published')
    lessons = course.lessons.all()
    
    # Check if student is enrolled
    enrollment = Enrollment.objects.filter(student=user, course=course).first()
    
    # Build lesson videos dict: lesson_id -> list of videos
    lesson_videos = {}
    for lesson in lessons:
        vids = list(lesson.videos.all())
        lesson_videos[lesson.id] = vids

    context = {
        'course': course,
        'lessons': lessons,
        'lesson_videos': lesson_videos,
        'enrollment': enrollment,
        'is_enrolled': enrollment is not None,
        'session_user': user,
        'page_title': course.title,
    }
    return render(request, 'student/courses/detail.html', context)


@session_login_required
@session_role_required('student')
def student_course_enroll(request, course_id):
    """Student can enroll in a course"""
    user = request.session_user
    course = get_object_or_404(Course, id=course_id, status='published')
    
    # Check if already enrolled
    if Enrollment.objects.filter(student=user, course=course).exists():
        messages.info(request, 'Vous êtes déjà inscrit à ce cours.')
        return redirect('student_course_detail', course_id=course.id)
    
    # Create enrollment
    enrollment = Enrollment(student=user, course=course, status='active')
    enrollment.save()
    
    # Create progress records for all lessons
    for lesson in course.lessons.all():
        Progress.objects.create(enrollment=enrollment, lesson=lesson, is_completed=False)
    
    messages.success(request, 'Inscription réussie !')
    return redirect('student_my_courses')


@session_login_required
@session_role_required('student')
def student_my_courses(request):
    """Student can view their enrolled courses"""
    user = request.session_user
    enrollments = Enrollment.objects.filter(student=user).select_related('course').order_by('-enrolled_at')
    
    context = {
        'enrollments': enrollments,
        'session_user': user,
        'page_title': 'Mes Cours',
    }
    return render(request, 'student/courses/my_courses.html', context)


@session_login_required
@session_role_required('student')
def student_course_learn(request, course_id):
    """Student can learn a course - view lessons and track progress"""
    user = request.session_user
    enrollment = get_object_or_404(Enrollment, student=user, course_id=course_id)
    course = enrollment.course
    lessons = course.lessons.all()
    
    # Get or create progress records
    progress_dict = {}
    for lesson in lessons:
        progress, _ = Progress.objects.get_or_create(enrollment=enrollment, lesson=lesson)
        progress_dict[lesson.id] = progress

    # Calculate overall progress (lesson-based)
    total_lessons = lessons.count()
    completed_lessons = enrollment.progress_records.filter(is_completed=True).count()
    progress_percentage = int((completed_lessons / total_lessons * 100)) if total_lessons > 0 else 0

    # Video-based progress
    total_videos = LessonVideo.objects.filter(lesson__course=course).count()
    completed_video_ids = set(VideoCompletion.objects.filter(
        user=user, lesson_video__lesson__course=course
    ).values_list('lesson_video_id', flat=True))
    video_progress_pct = int((len(completed_video_ids) / total_videos) * 100) if total_videos > 0 else 0

    # Build lesson videos dict: lesson_id -> list of videos
    lesson_videos = {}
    for lesson in lessons:
        vids = list(lesson.videos.all())
        lesson_videos[lesson.id] = vids

    # Use video progress if videos exist, otherwise lesson progress
    display_progress = video_progress_pct if total_videos > 0 else progress_percentage

    context = {
        'enrollment': enrollment,
        'course': course,
        'lessons': lessons,
        'progress_dict': progress_dict,
        'progress_percentage': display_progress,
        'total_lessons': total_lessons,
        'completed_lessons': completed_lessons,
        'total_videos': total_videos,
        'completed_videos': len(completed_video_ids),
        'completed_video_ids': completed_video_ids,
        'lesson_videos': lesson_videos,
        'session_user': user,
        'page_title': f'Apprendre - {course.title}',
    }
    return render(request, 'student/courses/learn.html', context)


@session_login_required
@session_role_required('student')
def student_lesson_complete(request, course_id, lesson_id):
    """Student can mark a lesson as complete"""
    user = request.session_user
    enrollment = get_object_or_404(Enrollment, student=user, course_id=course_id)
    lesson = get_object_or_404(Lesson, id=lesson_id, course_id=course_id)
    
    progress, created = Progress.objects.get_or_create(enrollment=enrollment, lesson=lesson)
    if not progress.is_completed:
        progress.is_completed = True
        progress.completed_at = timezone.now()
        progress.save()
        
        # Check if course is now 100% complete
        total_lessons = enrollment.course.lessons.count()
        completed_lessons = enrollment.progress_records.filter(is_completed=True).count()
        
        if completed_lessons == total_lessons and total_lessons > 0:
            # Mark enrollment as completed
            enrollment.status = 'completed'
            enrollment.completed_at = timezone.now()
            enrollment.save()
            
            # Generate certificate
            certificate_number = f"CERT-{uuid.uuid4().hex[:8].upper()}"
            Certificate.objects.create(
                enrollment=enrollment,
                certificate_number=certificate_number,
            )
            messages.success(request, 'Félicitations ! Vous avez terminé le cours et obtenu un certificat !')
    
    return redirect('student_course_learn', course_id=course_id)


@session_login_required
@session_role_required('student')
def student_video_complete(request):
    """AJAX endpoint: mark a video as completed for the current student"""
    import json as _json
    if request.method != 'POST':
        return JsonResponse({'error': 'POST only'}, status=405)
    user = request.session_user
    try:
        data = _json.loads(request.body)
        video_id = data.get('video_id')
    except Exception:
        return JsonResponse({'error': 'Invalid data'}, status=400)

    lesson_video = LessonVideo.objects.filter(id=video_id).select_related('lesson__course').first()
    if not lesson_video:
        return JsonResponse({'error': 'Video not found'}, status=404)

    # Verify enrollment
    course = lesson_video.lesson.course
    enrollment = Enrollment.objects.filter(student=user, course=course).first()
    if not enrollment:
        return JsonResponse({'error': 'Not enrolled'}, status=403)

    vc, created = VideoCompletion.objects.get_or_create(user=user, lesson_video=lesson_video)

    # Calculate video-based progress for this course
    total_videos = LessonVideo.objects.filter(lesson__course=course).count()
    completed_videos = VideoCompletion.objects.filter(
        user=user, lesson_video__lesson__course=course
    ).count()
    progress_pct = int((completed_videos / total_videos) * 100) if total_videos > 0 else 0

    # Auto-complete lesson if all its videos are done
    lesson = lesson_video.lesson
    lesson_video_ids = lesson.videos.values_list('id', flat=True)
    all_lesson_done = VideoCompletion.objects.filter(
        user=user, lesson_video_id__in=lesson_video_ids
    ).count() == lesson_video_ids.count() if lesson_video_ids else False

    course_completed = False
    if all_lesson_done:
        progress_obj, _ = Progress.objects.get_or_create(enrollment=enrollment, lesson=lesson)
        if not progress_obj.is_completed:
            progress_obj.is_completed = True
            progress_obj.completed_at = timezone.now()
            progress_obj.save()

        # Check full course completion
        total_lessons = course.lessons.count()
        completed_lessons = enrollment.progress_records.filter(is_completed=True).count()
        if completed_lessons == total_lessons and total_lessons > 0:
            enrollment.status = 'completed'
            enrollment.completed_at = timezone.now()
            enrollment.save()
            if not hasattr(enrollment, 'certificate') or not enrollment.certificate:
                Certificate.objects.create(
                    enrollment=enrollment,
                    certificate_number=f"CERT-{uuid.uuid4().hex[:8].upper()}",
                )
            course_completed = True

    return JsonResponse({
        'completed': True,
        'progress_percentage': progress_pct,
        'completed_videos': completed_videos,
        'total_videos': total_videos,
        'lesson_completed': all_lesson_done,
        'course_completed': course_completed,
    })


@session_login_required
@session_role_required('student')
def course_certificate_pdf(request, course_id):
    """Generate a professional PDF certificate for a completed course"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm, cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import Paragraph
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.pdfgen import canvas

    user = request.session_user
    enrollment = get_object_or_404(Enrollment, student=user, course_id=course_id, status='completed')

    # Verify 100% video completion
    total_videos = LessonVideo.objects.filter(lesson__course=enrollment.course).count()
    if total_videos > 0:
        completed_videos = VideoCompletion.objects.filter(
            user=user, lesson_video__lesson__course=enrollment.course
        ).count()
        if completed_videos < total_videos:
            return HttpResponse('Certificate only available at 100% completion', status=403)

    buf = io.BytesIO()
    page_w, page_h = landscape(A4)

    c = canvas.Canvas(buf, pagesize=landscape(A4))

    # ── Find certificate template image ──
    logo_path = os.path.join(settings.BASE_DIR, 'emsicertic.png')
    if not os.path.exists(logo_path):
        logo_path = os.path.join(settings.BASE_DIR, 'static', 'emsicertic.png')
    if not os.path.exists(logo_path):
        logo_path = os.path.join(settings.BASE_DIR, 'frontend', 'emsicertic.png')
    if not os.path.exists(logo_path):
        logo_path = None

    if logo_path:
        # ── Draw image as full-page background ──
        try:
            c.drawImage(logo_path, 0, 0, width=page_w, height=page_h, preserveAspectRatio=False)
        except Exception:
            c.setFillColor(colors.white)
            c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
    else:
        # ── Fallback: plain white background ──
        c.setFillColor(colors.white)
        c.rect(0, 0, page_w, page_h, fill=1, stroke=0)

    # ── Overlay dynamic text on top of the background image ──
    BLACK = colors.black
    margin = 18 * mm

    # Student full name — centered, large font, at the name line position
    student_name = f"{user.first_name} {user.last_name}".strip() or user.name
    c.setFont('Times-Bold', 28)
    c.setFillColor(BLACK)
    name_w = c.stringWidth(student_name, 'Times-Bold', 28)
    name_y = page_h * 0.48
    c.drawString((page_w - name_w) / 2, name_y, student_name)

    # Course / formation name — centered, at the training name line position
    course_name = enrollment.course.title.upper()
    c.setFont('Times-Bold', 14)
    c.setFillColor(BLACK)
    course_w = c.stringWidth(course_name, 'Times-Bold', 14)
    course_y = page_h * 0.35
    c.drawString((page_w - course_w) / 2, course_y, course_name)

    # Today's date — formatted dd/mm/yyyy, at the Date line bottom-left
    date_str = enrollment.completed_at.strftime('%d/%m/%Y') if enrollment.completed_at else ''
    c.setFont('Helvetica', 9)
    c.setFillColor(BLACK)
    c.drawString(margin + 20 * mm, margin + 18 * mm, date_str)

    c.save()
    buf.seek(0)

    # Save PDF to certificate record
    cert = getattr(enrollment, 'certificate', None)
    if cert:
        filename = f"cert_{enrollment.id}_{date_str.replace('/', '')}.pdf"
        cert.pdf_file.save(filename, buf, save=True)
        buf.seek(0)

    response = HttpResponse(buf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="certificate_{course_id}.pdf"'
    return response


@session_login_required
@session_role_required('student')
def student_my_certificates(request):
    """Student can view their earned certificates"""
    user = request.session_user
    certificates = Certificate.objects.filter(enrollment__student=user).select_related('enrollment__course').order_by('-issued_at')
    
    context = {
        'certificates': certificates,
        'session_user': user,
        'page_title': 'Mes Certificats',
    }
    return render(request, 'student/certificates/list.html', context)


@session_login_required
@session_role_required('student')
def student_certificate_view(request, certificate_id):
    """Student can view a specific certificate"""
    user = request.session_user
    certificate = get_object_or_404(Certificate, id=certificate_id, enrollment__student=user)
    
    context = {
        'certificate': certificate,
        'session_user': user,
        'page_title': f'Certificat {certificate.certificate_number}',
    }
    return render(request, 'student/certificates/view.html', context)


@session_login_required
@session_role_required('student')
def student_dashboard_updated(request):
    """Updated Student dashboard with courses, progress, and Student of the Month"""
    user = request.session_user
    
    # Get enrolled courses
    enrollments = Enrollment.objects.filter(student=user).select_related('course')[:6]
    
    # Get certificates
    certificates = Certificate.objects.filter(enrollment__student=user).select_related('enrollment__course')[:5]
    
    # Get current student of the month
    current_month = timezone.now().replace(day=1)
    student_of_month = StudentOfMonth.objects.filter(month=current_month).select_related('student').first()
    
    # Count stats
    enrolled_count = Enrollment.objects.filter(student=user).count()
    completed_count = Enrollment.objects.filter(student=user, status='completed').count()
    certificates_count = Certificate.objects.filter(enrollment__student=user).count()
    
    context = {
        'enrollments': enrollments,
        'certificates': certificates,
        'student_of_month': student_of_month,
        'is_student_of_month': student_of_month and student_of_month.student_id == user.id,
        'enrolled_count': enrolled_count,
        'completed_count': completed_count,
        'certificates_count': certificates_count,
        'session_user': user,
        'page_title': 'Dashboard Étudiant',
    }
    return render(request, 'student/dashboard.html', context)


# ── API Views for LMS ─────────────────────────────────────────────────────

@login_required_custom
@role_required('student')
@require_http_methods(['POST'])
def api_enroll_course(request, course_id):
    """API endpoint for student to enroll in a course"""
    user = request.custom_user
    
    try:
        course = Course.objects.get(id=course_id, status='published')
    except Course.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Course not found.'}, status=404)
    
    if Enrollment.objects.filter(student=user, course=course).exists():
        return JsonResponse({'success': False, 'message': 'Already enrolled.'}, status=400)
    
    enrollment = Enrollment(student=user, course=course, status='active')
    enrollment.save()
    
    # Create progress records
    for lesson in course.lessons.all():
        Progress.objects.create(enrollment=enrollment, lesson=lesson, is_completed=False)
    
    return JsonResponse({'success': True, 'message': 'Enrolled successfully.', 'enrollment_id': enrollment.id})


@login_required_custom
@role_required('student')
@require_http_methods(['POST'])
def api_complete_lesson(request, lesson_id):
    """API endpoint for student to mark a lesson as complete"""
    user = request.custom_user
    
    try:
        lesson = Lesson.objects.get(id=lesson_id)
        enrollment = Enrollment.objects.get(student=user, course=lesson.course)
    except (Lesson.DoesNotExist, Enrollment.DoesNotExist):
        return JsonResponse({'success': False, 'message': 'Lesson or enrollment not found.'}, status=404)
    
    progress, _ = Progress.objects.get_or_create(enrollment=enrollment, lesson=lesson)
    if not progress.is_completed:
        progress.is_completed = True
        progress.completed_at = timezone.now()
        progress.save()
        
        # Check if course completed
        total_lessons = lesson.course.lessons.count()
        completed_lessons = enrollment.progress_records.filter(is_completed=True).count()
        
        certificate_generated = False
        if completed_lessons == total_lessons and total_lessons > 0:
            enrollment.status = 'completed'
            enrollment.completed_at = timezone.now()
            enrollment.save()
            
            # Generate certificate if not exists
            if not hasattr(enrollment, 'certificate'):
                certificate_number = f"CERT-{uuid.uuid4().hex[:8].upper()}"
                Certificate.objects.create(
                    enrollment=enrollment,
                    certificate_number=certificate_number,
                )
                certificate_generated = True
        
        return JsonResponse({
            'success': True, 
            'message': 'Lesson marked as complete.',
            'progress_percentage': int((completed_lessons / total_lessons * 100)) if total_lessons > 0 else 0,
            'course_completed': certificate_generated,
        })
    
    return JsonResponse({'success': False, 'message': 'Lesson already completed.'}, status=400)


@login_required_custom
def api_get_progress(request, course_id):
    """API endpoint to get student's progress for a course"""
    user = request.custom_user
    
    try:
        enrollment = Enrollment.objects.get(student=user, course_id=course_id)
    except Enrollment.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Not enrolled.'}, status=404)
    
    total_lessons = enrollment.course.lessons.count()
    completed_lessons = enrollment.progress_records.filter(is_completed=True).count()
    progress_percentage = int((completed_lessons / total_lessons * 100)) if total_lessons > 0 else 0
    
    completed_lesson_ids = list(enrollment.progress_records.filter(is_completed=True).values_list('lesson_id', flat=True))
    
    return JsonResponse({
        'success': True,
        'total_lessons': total_lessons,
        'completed_lessons': completed_lessons,
        'progress_percentage': progress_percentage,
        'completed_lesson_ids': completed_lesson_ids,
        'enrollment_status': enrollment.status,
    })


# ── Student of the Month Logic ────────────────────────────────────────────

def calculate_student_of_the_month():
    """
    Calculate the Student of the Month based on completed courses.
    Should be run at the end of each month via cron job or management command.
    """
    now = timezone.now()
    first_day_of_month = now.replace(day=1)
    
    # Get all students who completed courses this month
    from django.db.models import Count
    
    completed_enrollments = Enrollment.objects.filter(
        status='completed',
        completed_at__year=now.year,
        completed_at__month=now.month,
    ).values('student').annotate(completed_count=Count('id')).order_by('-completed_count')
    
    if completed_enrollments:
        winner = completed_enrollments[0]
        student_id = winner['student']
        courses_completed = winner['completed_count']
        
        # Create or update Student of the Month record
        StudentOfMonth.objects.update_or_create(
            month=first_day_of_month,
            defaults={
                'student_id': student_id,
                'courses_completed_count': courses_completed,
            }
        )
        return student_id, courses_completed
    
    return None, 0


@session_login_required
def student_of_month(request):
    """Student of the Month page"""
    current_month = timezone.now().replace(day=1)
    student_of_month = StudentOfMonth.objects.filter(month=current_month).select_related('student').first()
    
    context = {
        'student_of_month': student_of_month,
        'session_user': request.session_user,
    }
    return render(request, 'student/student_of_month.html', context)


@session_login_required
@session_role_required('admin')
def admin_calculate_som(request):
    """Admin can manually trigger Student of the Month calculation"""
    student_id, courses_completed = calculate_student_of_the_month()
    
    if student_id:
        student = CustomUser.objects.get(id=student_id)
        messages.success(request, f'Étudiant du mois calculé : {student.name} avec {courses_completed} cours complétés.')
    else:
        messages.info(request, 'Aucun étudiant n\'a complété de cours ce mois-ci.')
    
    return redirect('admin_student_of_month')


@login_required_custom
@require_http_methods(['GET'])
def api_get_student_of_month(request):
    """API endpoint to get current Student of the Month"""
    current_month = timezone.now().replace(day=1)
    som = StudentOfMonth.objects.filter(month=current_month).select_related('student').first()
    
    if som:
        return JsonResponse({
            'success': True,
            'student': {
                'id': som.student.id,
                'name': som.student.name,
                'email': som.student.email,
            },
            'courses_completed': som.courses_completed_count,
            'month': som.month.strftime('%B %Y'),
        })
    
    return JsonResponse({'success': False, 'message': 'No student of the month yet.'}, status=404)


# ═══════════════════════════════════════════════════════════════════════
# CV MAKER VIEWS
# ═══════════════════════════════════════════════════════════════════════

import io
from .models import CV


def _generate_cv_pdf(cv):
    """Generate a polished, professional 2-column PDF from CV data using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, Image as RLImage, KeepTogether,
                                     Frame, PageTemplate, BaseDocTemplate)
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfgen import canvas as pdfcanvas

    buf = io.BytesIO()
    page_w, page_h = A4
    margin = 20 * mm
    usable = page_w - 2 * margin

    # ── Color palette ──
    C_TEXT = colors.HexColor('#1a1a1a')
    C_SECONDARY = colors.HexColor('#555555')
    C_RULE = colors.HexColor('#cccccc')
    C_ACCENT = colors.HexColor('#2c3e50')
    C_NAME_LAST = colors.HexColor('#777777')

    # ── Styles ──
    styles = getSampleStyleSheet()

    s_first = ParagraphStyle('CVFirst', parent=styles['Normal'], fontName='Helvetica-Bold',
                             fontSize=30, textColor=C_TEXT, spaceAfter=0, leading=34)
    s_last = ParagraphStyle('CVLast', parent=styles['Normal'], fontName='Helvetica-Bold',
                            fontSize=30, textColor=C_NAME_LAST, spaceAfter=0, leading=34)
    s_job = ParagraphStyle('CVJob', parent=styles['Normal'], fontName='Helvetica-Oblique',
                           fontSize=11, textColor=C_SECONDARY, spaceAfter=2*mm, leading=14)
    s_intro = ParagraphStyle('CVIntro', parent=styles['Normal'], fontName='Helvetica',
                             fontSize=9.5, textColor=C_SECONDARY, spaceAfter=0, leading=13)
    s_section = ParagraphStyle('CVSec', parent=styles['Normal'], fontName='Helvetica-Bold',
                               fontSize=9.5, textColor=C_ACCENT, spaceBefore=5*mm,
                               spaceAfter=0.8*mm, leading=12)
    s_entry_title = ParagraphStyle('CVET', parent=styles['Normal'], fontName='Helvetica-Bold',
                                   fontSize=10, textColor=C_TEXT, spaceAfter=0.3*mm, leading=13)
    s_entry_sub = ParagraphStyle('CVES', parent=styles['Normal'], fontName='Helvetica-Oblique',
                                 fontSize=9, textColor=C_SECONDARY, spaceAfter=0.3*mm, leading=12)
    s_bullet = ParagraphStyle('CVBul', parent=styles['Normal'], fontName='Helvetica',
                              fontSize=9, textColor=C_TEXT, leftIndent=8*mm,
                              bulletIndent=3*mm, spaceAfter=0.4*mm, leading=12)
    s_small = ParagraphStyle('CVSm', parent=styles['Normal'], fontName='Helvetica',
                             fontSize=9, textColor=C_TEXT, spaceAfter=0.5*mm, leading=12)
    s_skill = ParagraphStyle('CVSk', parent=styles['Normal'], fontName='Helvetica',
                             fontSize=9, textColor=C_TEXT, spaceAfter=0.3*mm, leading=12)

    # ── Extract data ──
    cv_data = cv.cv_data or {}
    sections = cv_data.get('sections', [])

    personal = {}
    other_sections = []
    for sec in sections:
        if sec.get('type') == 'personal_info':
            if sec.get('entries'):
                personal = sec['entries'][0] if sec['entries'] else {}
        else:
            other_sections.append(sec)

    first_name = personal.get('first_name') or cv.user.first_name
    last_name = personal.get('last_name') or cv.user.last_name
    job_title = personal.get('job_title', '')
    intro = personal.get('intro', '')
    phone = personal.get('phone') or cv.user.phone
    email = personal.get('email') or cv.user.email
    city = personal.get('city', '')
    country = personal.get('country', '')
    linkedin = personal.get('linkedin', '')

    # ── Classify sections ──
    LEFT_TYPES = {'skills', 'experience', 'custom'}
    RIGHT_TYPES = {'education', 'languages', 'interests', 'references'}

    left_sections = []
    right_sections = []

    # Contact block with unicode icons
    contact_lines = []
    if email:
        contact_lines.append(Paragraph(f"\u2709  {email}", s_small))
    if phone:
        contact_lines.append(Paragraph(f"\u260e  {phone}", s_small))
    addr = f"{city}{', ' + country if country else ''}" if city else ''
    if addr:
        contact_lines.append(Paragraph(f"\U0001f4cd  {addr}", s_small))
    if linkedin:
        contact_lines.append(Paragraph(f"\U0001f517  {linkedin}", s_small))
    if contact_lines:
        right_sections.insert(0, {'type': '_contact', 'title': 'CONTACT', 'entries': [], '_items': contact_lines})

    for sec in other_sections:
        sec_type = sec.get('type', 'custom')
        if sec_type in LEFT_TYPES:
            left_sections.append(sec)
        else:
            right_sections.append(sec)

    # ── Helper: render section ──
    def render_section(sec, width):
        items = []
        sec_type = sec.get('type', 'custom')
        title = sec.get('title', '')
        entries = sec.get('entries', [])

        # Section heading — UPPERCASE bold + thin colored underline
        items.append(Paragraph(title.upper(), s_section))
        line = Table([['']], colWidths=[width])
        line.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, C_RULE)]))
        items.append(line)
        items.append(Spacer(1, 1.5*mm))

        if sec_type == '_contact':
            items.insert(0, Spacer(1, 14))
            items.extend(sec.get('_items', []))
            return items

        if sec_type == 'skills':
            names = [e.get('name', '') for e in entries if e.get('name')]
            if names:
                mid = (len(names) + 1) // 2
                col1 = names[:mid]
                col2 = names[mid:]
                rows = []
                for i in range(max(len(col1), len(col2))):
                    rows.append([
                        Paragraph(f'\u2022 {col1[i]}' if i < len(col1) else '', s_skill),
                        Paragraph(f'\u2022 {col2[i]}' if i < len(col2) else '', s_skill),
                    ])
                t = Table(rows, colWidths=[width * 0.5, width * 0.5])
                t.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 2*mm),
                    ('TOPPADDING', (0, 0), (-1, -1), 0.5*mm),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 0.5*mm),
                ]))
                items.append(t)
            return items

        if sec_type == 'languages':
            for e in entries:
                name = e.get('name', '')
                level = e.get('level', '')
                if name:
                    txt = f"<b>{name}</b>" + (f" \u2014 {level}" if level else '')
                    items.append(Paragraph(txt, s_small))
            return items

        if sec_type == 'interests':
            names = [e.get('name', '') for e in entries if e.get('name')]
            if names:
                items.append(Paragraph(' \u2022 '.join(names), s_small))
            return items

        if sec_type == 'references':
            for e in entries:
                text = e.get('text', '')
                if text:
                    items.append(Paragraph(text, s_small))
            return items

        if sec_type == 'experience':
            for e in entries:
                title_val = e.get('title', '')
                company = e.get('company', '')
                period = e.get('period', '')
                location = e.get('location', '')
                bullets = e.get('bullets', '')

                if title_val or period:
                    # Job title left, date right-aligned on same line
                    if title_val and period:
                        row = Table([
                            [Paragraph(f"<b>{title_val}</b>", s_entry_title),
                             Paragraph(period, ParagraphStyle('CVDate', parent=s_entry_title,
                                      fontName='Helvetica', textColor=C_SECONDARY, alignment=TA_RIGHT))]
                        ], colWidths=[width * 0.68, width * 0.32])
                        row.setStyle(TableStyle([
                            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 0),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                            ('TOPPADDING', (0, 0), (-1, -1), 0),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                        ]))
                        items.append(row)
                    elif title_val:
                        items.append(Paragraph(f"<b>{title_val}</b>", s_entry_title))
                    else:
                        items.append(Paragraph(period, s_entry_title))

                if company or location:
                    sub = company
                    if location:
                        sub += f" \u2014 {location}" if sub else location
                    items.append(Paragraph(sub, s_entry_sub))

                if bullets:
                    for b in bullets.split('\n'):
                        b = b.strip()
                        if b:
                            prefix = '\u2022 ' if not b.startswith('\u2022') else ''
                            items.append(Paragraph(f"{prefix}{b}", s_bullet))
            return items

        if sec_type == 'education':
            for e in entries:
                title_val = e.get('title', '')
                institution = e.get('institution', '')
                period = e.get('period', '')
                location = e.get('location', '')

                if title_val:
                    items.append(Paragraph(f"<b>{title_val}</b>", s_entry_title))
                sub_parts = []
                if institution:
                    sub_parts.append(institution)
                if location:
                    sub_parts.append(location)
                if sub_parts:
                    items.append(Paragraph(' \u2014 '.join(sub_parts), s_entry_sub))
                if period:
                    items.append(Paragraph(period, ParagraphStyle('CVDat2', parent=s_small,
                               textColor=C_SECONDARY, fontSize=8.5)))
            return items

        # custom or fallback
        for e in entries:
            title_val = e.get('title', '')
            subtitle = e.get('subtitle', '')
            period = e.get('period', '')
            desc = e.get('description', '')

            if title_val:
                items.append(Paragraph(f"<b>{title_val}</b>", s_entry_title))
            if subtitle:
                items.append(Paragraph(subtitle, s_entry_sub))
            if desc:
                for b in desc.split('\n'):
                    b = b.strip()
                    if b:
                        prefix = '\u2022 ' if not b.startswith('\u2022') else ''
                        items.append(Paragraph(f"{prefix}{b}", s_bullet))
        return items

    # ── Build header ──
    header_elements = []

    # Name line: FIRST (bold dark) LAST (bold lighter gray)
    if cv.photo:
        try:
            img = RLImage(cv.photo.path, width=22*mm, height=22*mm)
            name_para = Paragraph(
                f'<font name="Helvetica-Bold" color="#1a1a1a">{first_name.upper()}</font>'
                f' <font name="Helvetica-Bold" color="#777777">{last_name.upper()}</font>',
                s_first)
            row = Table([[name_para, img]], colWidths=[usable - 25*mm, 25*mm])
            row.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            header_elements.append(row)
        except Exception:
            header_elements.append(Paragraph(
                f'<font name="Helvetica-Bold" color="#1a1a1a">{first_name.upper()}</font>'
                f' <font name="Helvetica-Bold" color="#777777">{last_name.upper()}</font>',
                s_first))
    else:
        header_elements.append(Paragraph(
            f'<font name="Helvetica-Bold" color="#1a1a1a">{first_name.upper()}</font>'
            f' <font name="Helvetica-Bold" color="#777777">{last_name.upper()}</font>',
            s_first))

    if job_title:
        header_elements.append(Paragraph(job_title, s_job))

    # Thin dark horizontal rule
    hr = Table([['']], colWidths=[usable])
    hr.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1.2, C_TEXT)]))
    header_elements.append(hr)
    header_elements.append(Spacer(1, 2*mm))

    if intro:
        header_elements.append(Paragraph(intro, s_intro))
        header_elements.append(Spacer(1, 3*mm))

    # ── Build column flowables ──
    left_w = usable * 0.62
    right_w = usable * 0.38
    col_gap = 4 * mm

    left_flow = []
    for sec in left_sections:
        left_flow.extend(render_section(sec, left_w - col_gap / 2))

    right_flow = []
    for sec in right_sections:
        right_flow.extend(render_section(sec, right_w - col_gap / 2))

    # ── Assemble story ──
    story = header_elements

    if left_flow or right_flow:
        max_len = max(len(left_flow), len(right_flow))
        while len(left_flow) < max_len:
            left_flow.append(Paragraph('', s_small))
        while len(right_flow) < max_len:
            right_flow.append(Paragraph('', s_small))

        rows = list(zip(left_flow, right_flow))
        col_table = Table(rows, colWidths=[left_w, right_w])
        col_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (0, -1), 0),
            ('RIGHTPADDING', (0, 0), (0, -1), col_gap / 2),
            ('LEFTPADDING', (1, 0), (1, -1), col_gap / 2),
            ('RIGHTPADDING', (1, 0), (1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(col_table)

    # ── Draw vertical separator via canvas callback ──
    def draw_vertical_line(canvas, doc):
        canvas.saveState()
        x = margin + left_w + col_gap / 2
        canvas.setStrokeColor(C_RULE)
        canvas.setLineWidth(0.5)
        canvas.line(x, margin, x, page_h - margin)
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=margin, bottomMargin=margin,
                            leftMargin=margin, rightMargin=margin)
    doc.build(story, onFirstPage=draw_vertical_line, onLaterPages=draw_vertical_line)
    buf.seek(0)
    return buf


def _generate_cv_docx(cv):
    """Generate a professional 2-column DOCX from CV data using python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    BLACK = RGBColor(0, 0, 0)
    GRAY = RGBColor(0x55, 0x55, 0x55)

    cv_data = cv.cv_data or {}
    sections = cv_data.get('sections', [])

    # Extract personal_info
    personal = {}
    other_sections = []
    for sec in sections:
        if sec.get('type') == 'personal_info':
            if sec.get('entries'):
                personal = sec['entries'][0] if sec['entries'] else {}
        else:
            other_sections.append(sec)

    first_name = personal.get('first_name') or cv.user.first_name
    last_name = personal.get('last_name') or cv.user.last_name
    job_title = personal.get('job_title', '')
    intro = personal.get('intro', '')
    phone = personal.get('phone') or cv.user.phone
    email = personal.get('email') or cv.user.email
    city = personal.get('city', '')
    country = personal.get('country', '')
    linkedin = personal.get('linkedin', '')

    # ── Header ──
    # Name
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(f"{first_name.upper()} {last_name.upper()}")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLACK

    # Job title
    if job_title:
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run(job_title)
        run.font.size = Pt(13)
        run.font.color.rgb = GRAY

    # Intro
    if intro:
        p = doc.add_paragraph()
        p.space_after = Pt(6)
        run = p.add_run(intro)
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY

    # Horizontal rule (thin border on empty paragraph)
    p = doc.add_paragraph()
    p.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Classify sections ──
    LEFT_TYPES = {'skills', 'experience', 'custom'}
    RIGHT_TYPES = {'education', 'languages', 'interests', 'references'}

    left_sections = []
    right_sections = []

    # Contact block for right column
    contact_items = []
    if phone:
        contact_items.append(f"Tél: {phone}")
    if email:
        contact_items.append(f"Email: {email}")
    if city:
        addr = city + (f", {country}" if country else '')
        contact_items.append(f"Adresse: {addr}")
    if linkedin:
        contact_items.append(f"LinkedIn: {linkedin}")
    if contact_items:
        right_sections.insert(0, {'type': '_contact', 'title': 'CONTACT', 'entries': [], '_items': contact_items})

    for sec in other_sections:
        sec_type = sec.get('type', 'custom')
        if sec_type in LEFT_TYPES:
            left_sections.append(sec)
        else:
            right_sections.append(sec)

    # ── Helper: add section heading ──
    def add_section_heading(doc_or_cell, title_text):
        p = doc_or_cell.add_paragraph()
        p.space_before = Pt(12)
        p.space_after = Pt(2)
        run = p.add_run(title_text.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
        # Underline border
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        btm = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '000000'
        })
        pBdr.append(btm)
        pPr.append(pBdr)

    # ── Helper: render section into a cell ──
    def render_section_in_cell(cell, sec):
        sec_type = sec.get('type', 'custom')
        title = sec.get('title', '')
        entries = sec.get('entries', [])

        add_section_heading(cell, title)

        if sec_type == '_contact':
            for item in sec.get('_items', []):
                p = cell.add_paragraph(item)
                p.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9)
            return

        if sec_type == 'skills':
            names = [e.get('name', '') for e in entries if e.get('name')]
            if names:
                mid = (len(names) + 1) // 2
                col1 = names[:mid]
                col2 = names[mid:]
                # Use a nested table for 2 sub-columns
                tbl = cell.add_table(rows=max(len(col1), len(col2)), cols=2)
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                for i in range(len(col1)):
                    c = tbl.cell(i, 0)
                    c.text = ''
                    p = c.paragraphs[0]
                    run = p.add_run(col1[i])
                    run.font.size = Pt(9)
                for i in range(len(col2)):
                    c = tbl.cell(i, 1)
                    c.text = ''
                    p = c.paragraphs[0]
                    run = p.add_run(col2[i])
                    run.font.size = Pt(9)
                # Remove borders from nested table
                for row in tbl.rows:
                    for c in row.cells:
                        tc = c._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            b = tcBorders.makeelement(qn(f'w:{border_name}'), {
                                qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                            })
                            tcBorders.append(b)
                        tcPr.append(tcBorders)
            return

        if sec_type == 'languages':
            for e in entries:
                name = e.get('name', '')
                level = e.get('level', '')
                if name:
                    p = cell.add_paragraph()
                    p.space_after = Pt(1)
                    run = p.add_run(name)
                    run.bold = True
                    run.font.size = Pt(9)
                    if level:
                        run2 = p.add_run(f" — {level}")
                        run2.font.size = Pt(9)
                        run2.font.color.rgb = GRAY
            return

        if sec_type == 'interests':
            names = [e.get('name', '') for e in entries if e.get('name')]
            if names:
                p = cell.add_paragraph(' • '.join(names))
                p.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9)
            return

        if sec_type == 'references':
            for e in entries:
                text = e.get('text', '')
                if text:
                    p = cell.add_paragraph(text)
                    p.space_after = Pt(1)
                    for run in p.runs:
                        run.font.size = Pt(9)
            return

        if sec_type == 'experience':
            for e in entries:
                title_val = e.get('title', '')
                company = e.get('company', '')
                period = e.get('period', '')
                location = e.get('location', '')
                bullets = e.get('bullets', '')

                if title_val:
                    p = cell.add_paragraph()
                    p.space_after = Pt(1)
                    run = p.add_run(title_val)
                    run.bold = True
                    run.font.size = Pt(10)
                    if period:
                        run2 = p.add_run(f"  |  {period}")
                        run2.font.size = Pt(10)
                        run2.font.color.rgb = GRAY
                if company or location:
                    sub = company
                    if location:
                        sub += f" — {location}" if sub else location
                    p = cell.add_paragraph(sub)
                    p.space_after = Pt(1)
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = GRAY
                if bullets:
                    for b in bullets.split('\n'):
                        b = b.strip()
                        if b:
                            prefix = '• ' if not b.startswith('•') else ''
                            p = cell.add_paragraph(f"{prefix}{b}")
                            p.space_after = Pt(0)
                            p.paragraph_format.left_indent = Cm(0.5)
                            for run in p.runs:
                                run.font.size = Pt(9)
            return

        if sec_type == 'education':
            for e in entries:
                title_val = e.get('title', '')
                institution = e.get('institution', '')
                period = e.get('period', '')
                location = e.get('location', '')

                if title_val:
                    p = cell.add_paragraph()
                    p.space_after = Pt(1)
                    run = p.add_run(title_val)
                    run.bold = True
                    run.font.size = Pt(10)
                    if period:
                        run2 = p.add_run(f"  |  {period}")
                        run2.font.size = Pt(10)
                        run2.font.color.rgb = GRAY
                if institution or location:
                    sub = institution
                    if location:
                        sub += f" — {location}" if sub else location
                    p = cell.add_paragraph(sub)
                    p.space_after = Pt(1)
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = GRAY
            return

        # custom or fallback
        for e in entries:
            title_val = e.get('title', '')
            subtitle = e.get('subtitle', '')
            period = e.get('period', '')
            desc = e.get('description', '')

            if title_val:
                p = cell.add_paragraph()
                p.space_after = Pt(1)
                run = p.add_run(title_val)
                run.bold = True
                run.font.size = Pt(10)
                if period:
                    run2 = p.add_run(f"  |  {period}")
                    run2.font.size = Pt(10)
                    run2.font.color.rgb = GRAY
            if subtitle:
                p = cell.add_paragraph(subtitle)
                p.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = GRAY
            if desc:
                for b in desc.split('\n'):
                    b = b.strip()
                    if b:
                        prefix = '• ' if not b.startswith('•') else ''
                        p = cell.add_paragraph(f"{prefix}{b}")
                        p.space_after = Pt(0)
                        p.paragraph_format.left_indent = Cm(0.5)
                        for run in p.runs:
                            run.font.size = Pt(9)

    # ── Build 2-column table ──
    if left_sections or right_sections:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_cell = tbl.cell(0, 0)
        right_cell = tbl.cell(0, 1)

        # Set column widths (63% / 37%)
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        left_cell.width = int(page_width * 0.63)
        right_cell.width = int(page_width * 0.37)

        # Add vertical border between columns
        # Set right border on left cell
        tc = left_cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        right_bdr = tcBorders.makeelement(qn('w:right'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '0', qn('w:color'): '555555'
        })
        tcBorders.append(right_bdr)
        tcPr.append(tcBorders)

        # Remove all other borders from both cells
        for cell in [left_cell, right_cell]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
            for bn in ['top', 'left', 'bottom']:
                if cell == right_cell and bn == 'left':
                    continue  # skip left border of right cell (handled by left cell's right border)
                b = tcBorders.makeelement(qn(f'w:{bn}'), {
                    qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                })
                tcBorders.append(b)
            tcPr.append(tcBorders)

        # Render sections
        for sec in left_sections:
            render_section_in_cell(left_cell, sec)
        for sec in right_sections:
            render_section_in_cell(right_cell, sec)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@session_login_required
def cv_list(request):
    """List the current user's CVs."""
    user = request.session_user
    cvs = CV.objects.filter(user=user)
    return render(request, 'dashboard/cv/list.html', {
        'session_user': user,
        'cvs': cvs,
        'page_title': 'Mes CVs',
    })


@session_login_required
def cv_create(request):
    """Create a new CV."""
    user = request.session_user
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        cv_data_json = request.POST.get('cv_data', '{}')
        photo = request.FILES.get('photo')

        if not title:
            messages.error(request, 'Le titre est requis.')
            return redirect('cv_create')

        try:
            cv_data = json.loads(cv_data_json)
        except json.JSONDecodeError:
            cv_data = {}

        cv = CV(user=user, title=title, cv_data=cv_data)
        if photo:
            cv.photo = photo
        cv.save()
        messages.success(request, 'CV créé avec succès.')
        return redirect('cv_list')

    return render(request, 'dashboard/cv/create.html', {
        'session_user': user,
        'page_title': 'Créer un CV',
        'cv': None,
        'cv_data_json': '{}',
    })


@session_login_required
def cv_edit(request, id):
    """Edit an existing CV."""
    user = request.session_user
    cv = get_object_or_404(CV, id=id, user=user)

    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        cv_data_json = request.POST.get('cv_data', '{}')
        photo = request.FILES.get('photo')
        remove_photo = request.POST.get('remove_photo') == '1'

        if not title:
            messages.error(request, 'Le titre est requis.')
            return redirect('cv_edit', id=id)

        try:
            cv_data = json.loads(cv_data_json)
        except json.JSONDecodeError:
            cv_data = {}

        cv.title = title
        cv.cv_data = cv_data
        if photo:
            cv.photo = photo
        elif remove_photo and cv.photo:
            cv.photo.delete(save=False)
            cv.photo = None
        cv.save()
        messages.success(request, 'CV mis à jour avec succès.')
        return redirect('cv_list')

    cv_data_json = json.dumps(cv.cv_data)
    return render(request, 'dashboard/cv/create.html', {
        'session_user': user,
        'page_title': 'Modifier le CV',
        'cv': cv,
        'cv_data_json': cv_data_json,
    })


@session_login_required
@require_http_methods(['POST'])
def cv_delete(request, id):
    """Delete a CV (POST only)."""
    user = request.session_user
    cv = get_object_or_404(CV, id=id, user=user)
    cv.delete()
    messages.success(request, 'CV supprimé.')
    return redirect('cv_list')


@session_login_required
@require_http_methods(['POST'])
def cv_clone(request, id):
    """Clone a CV (POST only)."""
    user = request.session_user
    cv = get_object_or_404(CV, id=id, user=user)
    new_cv = CV(
        user=user,
        title=f"{cv.title} (copie)",
        cv_data=cv.cv_data,
    )
    new_cv.save()
    messages.success(request, 'CV cloné avec succès.')
    return redirect('cv_list')


@session_login_required
def cv_download_pdf(request, id):
    """Download CV as PDF."""
    user = request.session_user
    cv = get_object_or_404(CV, id=id, user=user)
    buf = _generate_cv_pdf(cv)
    response = HttpResponse(buf, content_type='application/pdf')
    filename = f"{cv.title.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@session_login_required
def cv_download_docx(request, id):
    """Download CV as DOCX."""
    user = request.session_user
    cv = get_object_or_404(CV, id=id, user=user)
    buf = _generate_cv_docx(cv)
    response = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"{cv.title.replace(' ', '_')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── Chat Public Views ─────────────────────────────────────────────────────

@session_login_required
def chat_page(request):
    """Render the public chat page."""
    user = request.session_user
    is_banned = ChatBan.objects.filter(user=user).exists()
    return render(request, 'dashboard/chat.html', {
        'session_user': user,
        'is_banned': is_banned,
        'is_admin': user.role == 'admin',
    })


@session_login_required
@require_http_methods(['POST'])
def chat_send(request):
    """Send a chat message. Banned users cannot send."""
    user = request.session_user
    is_banned = ChatBan.objects.filter(user=user).exists()
    if is_banned:
        return JsonResponse({'success': False, 'message': 'Vous avez été désactivé du chat'}, status=403)
    data = _parse_json(request)
    message = (data.get('message') or '').strip()
    if not message:
        return JsonResponse({'success': False, 'message': 'Message vide.'}, status=400)
    ChatMessage.objects.create(user=user, message=message)
    return JsonResponse({'success': True})


@session_login_required
@require_http_methods(['GET'])
def chat_messages(request):
    """Return last 100 non-deleted messages as JSON."""
    user = request.session_user
    msgs = ChatMessage.objects.filter(is_deleted=False).select_related('user').order_by('-created_at')[:100]
    msgs = reversed(list(msgs))

    # Precompute friendship info for the current user
    friend_ids = set()
    for f in Friendship.objects.filter(Q(user1=user) | Q(user2=user)):
        friend_ids.add(f.user1_id if f.user2_id == user.id else f.user2_id)
    pending_sent_ids = set(FriendRequest.objects.filter(from_user=user, status='pending').values_list('to_user_id', flat=True))
    blocked_ids = set(BlockedUser.objects.filter(blocker=user).values_list('blocked_id', flat=True))

    data = []
    for m in msgs:
        # Determine friendship status
        if m.user.id == user.id:
            friendship_status = 'self'
        elif m.user.id in blocked_ids:
            friendship_status = 'blocked'
        elif m.user.id in friend_ids:
            friendship_status = 'friend'
        elif m.user.id in pending_sent_ids:
            friendship_status = 'pending'
        else:
            friendship_status = 'none'

        data.append({
            'id': m.id,
            'username': m.user.name,
            'message': m.message,
            'created_at': m.created_at.strftime('%d/%m/%Y %H:%M'),
            'is_admin': m.user.role == 'admin',
            'is_own': m.user.id == user.id,
            'user_id': m.user.id,
            'is_banned': ChatBan.objects.filter(user=m.user).exists(),
            'friendship_status': friendship_status,
        })
    return JsonResponse({'success': True, 'messages': data})


@session_login_required
@require_http_methods(['POST'])
def chat_delete(request, message_id):
    """Soft-delete a message. Admin can delete any, user only own."""
    user = request.session_user
    msg = get_object_or_404(ChatMessage, id=message_id, is_deleted=False)
    if user.role != 'admin' and msg.user.id != user.id:
        return JsonResponse({'success': False, 'message': 'Non autorisé.'}, status=403)
    msg.is_deleted = True
    msg.save()
    return JsonResponse({'success': True})


@session_login_required
@require_http_methods(['POST'])
def chat_ban(request, user_id):
    """Admin only: ban a user from chat."""
    admin = request.session_user
    if admin.role != 'admin':
        return JsonResponse({'success': False, 'message': 'Admin uniquement.'}, status=403)
    target = get_object_or_404(CustomUser, id=user_id)
    if target.role == 'admin':
        return JsonResponse({'success': False, 'message': 'Cannot ban admin.'}, status=403)
    ChatBan.objects.get_or_create(user=target, defaults={'banned_by': admin})
    return JsonResponse({'success': True})


@session_login_required
@require_http_methods(['POST'])
def chat_unban(request, user_id):
    """Admin only: unban a user from chat."""
    admin = request.session_user
    if admin.role != 'admin':
        return JsonResponse({'success': False, 'message': 'Admin uniquement.'}, status=403)
    target = get_object_or_404(CustomUser, id=user_id)
    ChatBan.objects.filter(user=target).delete()
    return JsonResponse({'success': True})


# ── Friend System & Private Messaging Views ───────────────────────────────

@session_login_required
@require_http_methods(['POST'])
def chat_friend_request(request, user_id):
    """Send a friend request to another user."""
    user = request.session_user
    target = get_object_or_404(CustomUser, id=user_id)
    if target.id == user.id:
        return JsonResponse({'success': False, 'message': 'Cannot friend yourself.'}, status=400)
    # Check if blocked
    if BlockedUser.objects.filter(blocker=target, blocked=user).exists() or BlockedUser.objects.filter(blocker=user, blocked=target).exists():
        return JsonResponse({'success': False, 'message': 'Utilisateur bloqué.'}, status=400)
    # Check if already friends
    if Friendship.objects.filter(Q(user1_id=min(user.id, target.id), user2_id=max(user.id, target.id))).exists():
        return JsonResponse({'success': False, 'message': 'Déjà amis.'}, status=400)
    # Check if pending request already exists (either direction)
    if FriendRequest.objects.filter(from_user=user, to_user=target, status='pending').exists():
        return JsonResponse({'success': False, 'message': 'Demande déjà envoyée.'}, status=400)
    if FriendRequest.objects.filter(from_user=target, to_user=user, status='pending').exists():
        return JsonResponse({'success': False, 'message': 'Cet utilisateur vous a déjà envoyé une demande.'}, status=400)
    FriendRequest.objects.create(from_user=user, to_user=target)
    return JsonResponse({'success': True})


@session_login_required
@require_http_methods(['POST'])
def chat_friend_accept(request, request_id):
    """Accept a friend request."""
    user = request.session_user
    fr = get_object_or_404(FriendRequest, id=request_id, to_user=user, status='pending')
    fr.status = 'accepted'
    fr.save()
    # Create Friendship with user1_id < user2_id
    u1, u2 = (fr.from_user_id, fr.to_user_id) if fr.from_user_id < fr.to_user_id else (fr.to_user_id, fr.from_user_id)
    Friendship.objects.get_or_create(user1_id=u1, user2_id=u2)
    return JsonResponse({'success': True})


@session_login_required
@require_http_methods(['POST'])
def chat_friend_decline(request, request_id):
    """Decline a friend request."""
    user = request.session_user
    fr = get_object_or_404(FriendRequest, id=request_id, to_user=user, status='pending')
    fr.status = 'declined'
    fr.save()
    return JsonResponse({'success': True})


@session_login_required
def chat_invitations(request):
    """Show pending friend requests received by the logged-in user."""
    user = request.session_user
    pending_requests = FriendRequest.objects.filter(to_user=user, status='pending').select_related('from_user')
    return render(request, 'dashboard/invitations.html', {
        'session_user': user,
        'pending_requests': pending_requests,
    })


@session_login_required
@require_http_methods(['POST'])
def chat_block_user(request, user_id):
    """Block a user — also removes any existing friendship."""
    user = request.session_user
    target = get_object_or_404(CustomUser, id=user_id)
    if target.id == user.id:
        return JsonResponse({'success': False, 'message': 'Cannot block yourself.'}, status=400)
    BlockedUser.objects.get_or_create(blocker=user, blocked=target)
    # Remove friendship if exists
    u1, u2 = (user.id, target.id) if user.id < target.id else (target.id, user.id)
    Friendship.objects.filter(user1_id=u1, user2_id=u2).delete()
    return JsonResponse({'success': True})


@session_login_required
def chat_private(request, user_id):
    """Private message conversation with a specific user."""
    user = request.session_user
    other = get_object_or_404(CustomUser, id=user_id)
    if other.id == user.id:
        return redirect('chat_page')
    # Check if blocked
    if BlockedUser.objects.filter(Q(blocker=user, blocked=other) | Q(blocker=other, blocked=user)).exists():
        return redirect('chat_page')

    if request.method == 'POST':
        data = _parse_json(request)
        message = (data.get('message') or '').strip()
        if not message:
            return JsonResponse({'success': False, 'message': 'Message vide.'}, status=400)
        PrivateMessage.objects.create(sender=user, receiver=other, message=message)
        return JsonResponse({'success': True})

    # GET: mark all received messages from other as read
    PrivateMessage.objects.filter(sender=other, receiver=user, is_read=False).update(is_read=True)

    # Get conversation
    conversation = PrivateMessage.objects.filter(
        Q(sender=user, receiver=other) | Q(sender=other, receiver=user)
    ).select_related('sender', 'receiver').order_by('created_at')

    return render(request, 'dashboard/private_chat.html', {
        'session_user': user,
        'other_user': other,
        'conversation': conversation,
    })


@session_login_required
@require_http_methods(['GET'])
def chat_friends_json(request):
    """Return current user's friends list as JSON."""
    user = request.session_user
    friendships = Friendship.objects.filter(Q(user1=user) | Q(user2=user)).select_related('user1', 'user2')
    friends = []
    for f in friendships:
        friend = f.user2 if f.user1_id == user.id else f.user1
        # Get last private message
        last_msg = PrivateMessage.objects.filter(
            Q(sender=user, receiver=friend) | Q(sender=friend, receiver=user)
        ).order_by('-created_at').first()
        unread_count = PrivateMessage.objects.filter(sender=friend, receiver=user, is_read=False).count()
        friends.append({
            'id': friend.id,
            'name': friend.name,
            'role': friend.role,
            'unread_count': unread_count,
            'last_message': last_msg.message[:50] if last_msg else None,
            'last_message_time': last_msg.created_at.strftime('%d/%m/%Y %H:%M') if last_msg else None,
        })
    return JsonResponse({'success': True, 'friends': friends})
